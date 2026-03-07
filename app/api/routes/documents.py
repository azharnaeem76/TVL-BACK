"""Document Upload & AI Analysis API."""
import os
import uuid
import logging
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import Optional
from pydantic import BaseModel

from app.core.database import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.documents import Document, DocumentStatus

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/documents", tags=["Documents"])

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "uploads")
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_name: str
    file_type: str
    file_size: int
    status: DocumentStatus
    title: Optional[str] = None
    summary: Optional[str] = None
    extracted_parties: Optional[str] = None
    extracted_sections: Optional[str] = None
    extracted_court: Optional[str] = None
    extracted_judge: Optional[str] = None
    extracted_date: Optional[str] = None
    key_points: Optional[str] = None
    created_at: Optional[datetime] = None

    class Config:
        from_attributes = True


@router.post(
    "/upload",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Upload a legal document for AI analysis",
)
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    # Validate extension
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read and validate size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum 10MB allowed.")

    # Save file
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    unique_name = f"{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_name)
    with open(file_path, "wb") as f:
        f.write(content)

    # Create DB record
    doc = Document(
        user_id=user.id,
        filename=unique_name,
        original_name=file.filename or "unknown",
        file_type=ext.lstrip("."),
        file_size=len(content),
        status=DocumentStatus.UPLOADED,
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    return DocumentResponse.model_validate(doc)


@router.get("/", response_model=dict, summary="List user's uploaded documents")
async def list_documents(
    status_filter: Optional[DocumentStatus] = Query(None, alias="status"),
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    base = select(Document).where(Document.user_id == user.id)
    count_q = select(func.count(Document.id)).where(Document.user_id == user.id)

    if status_filter:
        base = base.where(Document.status == status_filter)
        count_q = count_q.where(Document.status == status_filter)

    total = (await db.execute(count_q)).scalar() or 0
    rows = (await db.execute(
        base.order_by(Document.created_at.desc()).offset(skip).limit(limit)
    )).scalars().all()

    return {
        "items": [DocumentResponse.model_validate(d) for d in rows],
        "total": total,
    }


@router.get("/{doc_id}", response_model=DocumentResponse, summary="Get document details")
async def get_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return DocumentResponse.model_validate(doc)


@router.post("/{doc_id}/analyze", response_model=DocumentResponse, summary="Trigger AI analysis")
async def analyze_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.status == DocumentStatus.ANALYZING:
        raise HTTPException(status_code=400, detail="Document is already being analyzed")

    doc.status = DocumentStatus.ANALYZING
    await db.flush()

    # Extract text from document
    file_path = os.path.join(UPLOAD_DIR, doc.filename)
    extracted_text = ""

    try:
        if doc.file_type == "pdf":
            try:
                import fitz  # PyMuPDF
                pdf_doc = fitz.open(file_path)
                extracted_text = "\n".join(page.get_text() for page in pdf_doc)
                pdf_doc.close()
            except ImportError:
                extracted_text = "[PDF text extraction requires PyMuPDF. Install: pip install PyMuPDF]"
        elif doc.file_type == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_text = f.read()
        elif doc.file_type in ("doc", "docx"):
            try:
                import docx
                word_doc = docx.Document(file_path)
                extracted_text = "\n".join(p.text for p in word_doc.paragraphs)
            except ImportError:
                extracted_text = "[DOCX extraction requires python-docx. Install: pip install python-docx]"

        if not extracted_text or extracted_text.startswith("["):
            # Fallback: basic analysis without AI
            doc.status = DocumentStatus.COMPLETED
            doc.summary = extracted_text if extracted_text.startswith("[") else "Text extraction produced no content."
            doc.title = doc.original_name
            await db.flush()
            await db.refresh(doc)
            return DocumentResponse.model_validate(doc)

        # Try AI analysis via Ollama
        try:
            import httpx
            from app.core.config import get_settings
            settings = get_settings()

            prompt = f"""Analyze this Pakistani legal document and extract the following in JSON format:
{{
  "title": "case or document title",
  "summary": "brief 2-3 sentence summary",
  "parties": "petitioner vs respondent or parties involved",
  "sections": "comma-separated list of legal sections cited (e.g., Section 302 PPC, Section 497 CrPC)",
  "court": "which court (e.g., Supreme Court, Lahore High Court)",
  "judge": "judge name(s) if mentioned",
  "date": "date of judgment if mentioned",
  "key_points": "3-5 key legal points, separated by semicolons"
}}

Document text (first 4000 chars):
{extracted_text[:4000]}"""

            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(
                    f"{settings.OLLAMA_BASE_URL}/api/generate",
                    json={"model": settings.OLLAMA_MODEL, "prompt": prompt, "stream": False},
                )

            if resp.status_code == 200:
                import json
                ai_text = resp.json().get("response", "")
                # Try to parse JSON from AI response
                try:
                    # Find JSON in response
                    start = ai_text.find("{")
                    end = ai_text.rfind("}") + 1
                    if start >= 0 and end > start:
                        parsed = json.loads(ai_text[start:end])
                        doc.title = parsed.get("title", doc.original_name)
                        doc.summary = parsed.get("summary", "")
                        doc.extracted_parties = parsed.get("parties", "")
                        doc.extracted_sections = parsed.get("sections", "")
                        doc.extracted_court = parsed.get("court", "")
                        doc.extracted_judge = parsed.get("judge", "")
                        doc.extracted_date = parsed.get("date", "")
                        doc.key_points = parsed.get("key_points", "")
                except (json.JSONDecodeError, KeyError):
                    doc.summary = ai_text[:2000]
                    doc.title = doc.original_name

            doc.status = DocumentStatus.COMPLETED

        except Exception as e:
            logger.warning(f"AI analysis failed: {e}")
            # Fallback: store raw text summary
            doc.status = DocumentStatus.COMPLETED
            doc.title = doc.original_name
            doc.summary = extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else "")

    except Exception as e:
        logger.error(f"Document analysis error: {e}")
        doc.status = DocumentStatus.FAILED
        doc.summary = f"Analysis failed: {str(e)}"

    await db.flush()
    await db.refresh(doc)
    return DocumentResponse.model_validate(doc)


@router.delete("/{doc_id}", status_code=status.HTTP_204_NO_CONTENT, summary="Delete a document")
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file
    file_path = os.path.join(UPLOAD_DIR, doc.filename)
    if os.path.exists(file_path):
        os.remove(file_path)

    await db.delete(doc)
    await db.flush()
